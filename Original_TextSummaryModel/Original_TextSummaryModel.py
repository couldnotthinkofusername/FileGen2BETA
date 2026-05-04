import docx
import PyPDF2 as pdf
import pdfplumber as pdl
import logging
import datetime
import re
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from collections import defaultdict
import math
import os

class TextExtraction:
    @staticmethod
    def DocumentReader(file_path):
        try:
            doc = docx.Document(file_path)
            text = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            text.append(cell_text)
            

            properties = doc.core_properties
            metadata = {
                'title': properties.title or 'Unknown',
                'author': properties.author or 'Unknown',
                'created': properties.created or datetime.datetime.now(),
                'modified': properties.modified or datetime.datetime.now(),
                'keywords': properties.keywords or '',
                'extraction_date': datetime.datetime.now().strftime("%Y-%m-%d"),
                'word_count': sum(len(word_tokenize(para)) for para in text)
            }
            
            cleaned_text = CleanText.clean(' '.join(text))
            return cleaned_text, metadata
        except Exception as e:
            logging.error(f"Error reading DOCX: {str(e)}")
            raise

    @staticmethod
    def PdfReaderWithPyPDF2(file_path):
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pdf.PdfReader(file)
                text = [page.extract_text() or '' for page in pdf_reader.pages]
                
                try:
                    with pdl.open(file_path) as pdf:
                        for page in pdf.pages:
                            tables = page.extract_tables()
                            for table in tables:
                                for row in table:
                                    row_text = ' '.join(str(cell) for cell in row if cell).strip()
                                    if row_text:
                                        text.append(row_text)
                except Exception as table_error:
                    logging.warning(f"Table extraction with pdfplumber failed: {str(table_error)}")
                
                metadata = pdf_reader.metadata or {}
                metadata.update({
                    'extraction_date': datetime.datetime.now().strftime("%Y-%m-%d"),
                    'word_count': sum(len(word_tokenize(page)) for page in text)
                })
                
                cleaned_text = CleanText.clean(' '.join(text))
                return cleaned_text, metadata
        except Exception as e:
            logging.error(f"Error reading PDF with PyPDF2: {str(e)}")
            raise

    @staticmethod
    def PdfReaderWithPlumber(file_path):
        try:
            with pdl.open(file_path) as pdf:
                text = []
                for page in pdf.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text.extend(line.strip() for line in extracted.splitlines() if line.strip())
                    
                
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            row_text = ' '.join(str(cell) for cell in row if cell).strip()
                            if row_text:
                                text.append(row_text)
                
                metadata = pdf.metadata or {}
                metadata.update({
                    'extraction_date': datetime.datetime.now().strftime("%Y-%m-%d"),
                    'word_count': sum(len(word_tokenize(line)) for line in text)
                })
                
                cleaned_text = CleanText.clean(' '.join(text))
                return cleaned_text, metadata
        except Exception as e:
            logging.error(f"Error reading PDF with pdfplumber: {str(e)}")
            raise

class CleanText:
    @staticmethod
    def clean(text):
        if not isinstance(text, str) or not text.strip():
            return ''
        
      
        text = re.sub(r'\s+', ' ', text)  # Normalize whitespace
        text = re.sub(r'[^\x20-\x7E]', '', text)  # Remove non-printable characters
        text = re.sub(r'[-–—]+', '-', text)  # Normalize dashes
        text = re.sub(r'[•●■◦♦]', '', text)  # Remove bullets and symbols
        text = re.sub(r'\bPage\s*\d+\b', '', text, flags=re.IGNORECASE)  # Remove page numbers
        text = re.sub(r'\b\d{1,2}[-/\.]\d{1,2}[-/\.]\d{2,4}\b', '', text)  # Remove dates
        text = re.sub(r'\s*([.,!?])\s*', r'\1 ', text)  # Normalize punctuation
        text = re.sub(r'\b(figure|table|section)\s*\d+\b', '', text, flags=re.IGNORECASE)  # Remove figure/table refs
        
        try:
            sentences = sent_tokenize(text)
            sentences = [s.strip() for s in sentences if len(word_tokenize(s)) > 4]
            text = ' '.join(s.rstrip('!?.') + '.' for s in sentences if s)
        except Exception as e:
            logging.error(f"Sentence tokenization error: {str(e)}")
            text = text.rstrip('!?.') + '.' if text else ''
        
        return text.strip()

class SummaryModel:
    @staticmethod
    def word_frequencies(text):
        stop_words = set(stopwords.words('english'))
        words = word_tokenize(text.lower())
        word_freq = defaultdict(int)
        
        for word in words:
            if word.isalnum() and word not in stop_words and len(word) > 3:
                word_freq[word] += 1
        
        return word_freq if word_freq else None

    @staticmethod
    def sentence_scores(sentences, word_freq):
        scores = {}
        for i, sentence in enumerate(sentences):
            words = word_tokenize(sentence.lower())
            word_count = sum(1 for word in words if word.isalnum())
            
            if word_count < 5: 
                continue
                
            score = sum(word_freq.get(word, 0) for word in words if word in word_freq) / (word_count + 1)
            scores[i] = score
            
        return scores

    @staticmethod
    def summarize(text, ratio=0.2, min_sentences=2, max_sentences=5):
        try:
            sentences = sent_tokenize(text)
            if not sentences:
                return '', {'error': 'No valid sentences found'}
                
            word_freq = SummaryModel.word_frequencies(text)
            if not word_freq:
                return '', {'error': 'No valid word frequencies found'}
                
            scores = SummaryModel.sentence_scores(sentences, word_freq)
            if not scores:
                return '', {'error': 'No scorable sentences found'}
                
            
            num_sentences = min(max_sentences, max(min_sentences, math.ceil(len(sentences) * ratio)))
            top_sentences = sorted(scores.items(), key=lambda x: x[1], reverse=True)[:num_sentences]
            
            summary_sentences = [sentences[i] for i, _ in sorted(top_sentences, key=lambda x: x[0])]
            summary = ' '.join(summary_sentences)
            
            return summary, {
                'original_sentences': len(sentences),
                'summary_sentences': len(summary_sentences),
                'compression_ratio': ratio
            }
        except Exception as e:
            logging.error(f"Summarization error: {str(e)}")
            return '', {'error': str(e)}

class ProcessorPipeline:
    @staticmethod
    def process_document(file_path, summarize=True, ratio=0.2):
        if not os.path.exists(file_path):
            return {'error': 'File does not exist'}
            
        file_extension = os.path.splitext(file_path)[1].lower()
        if file_extension not in ('.docx', '.pdf'):
            return {'error': 'Unsupported file type'}
            
        try:
            if file_extension == '.pdf':
                file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
                if file_size_mb > 10:
                    text, metadata = TextExtraction.PdfReaderWithPlumber(file_path)
                else:
                    text, metadata = TextExtraction.PdfReaderWithPyPDF2(file_path)
            else:
                text, metadata = TextExtraction.DocumentReader(file_path)
                
            if not text:
                return {'error': 'No text extracted'}
                
            result = {
                'text': text,
                'metadata': metadata
            }
            
            if summarize:
                summary, summary_metadata = SummaryModel.summarize(text, ratio)
                result['summary'] = summary
                result['summary_metadata'] = summary_metadata
                
            return result
            
        except Exception as e:
            logging.error(f"Processing error: {str(e)}")
            return {'error': str(e)}

if __name__ == '__main__':
    logging.basicConfig(level=logging.INFO)
    file_path = input('Enter path to DOCX or PDF file: ').strip()
    summarize = input('Generate summary? (yes/no): ').strip().lower() == 'yes'
    ratio = float(input('Enter summary ratio (0.1-0.5, default 0.3): ') or 0.3)
    
    result = ProcessorPipeline.process_document(file_path, summarize, ratio)
    print("\nResult:")
    for key, value in result.items():
        print(f"{key.capitalize()}: {value}")
